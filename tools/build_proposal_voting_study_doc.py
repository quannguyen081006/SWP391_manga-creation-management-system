from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
MD_PATH = OUT_DIR / "ON_TAP_MODULE_PROPOSAL_VOTING.md"
DOCX_PATH = OUT_DIR / "ON_TAP_MODULE_PROPOSAL_VOTING.docx"


def table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        lines.append("| " + " | ".join(str(c).replace("\n", "<br>") for c in row) + " |")
    return lines


def build_markdown():
    lines = []
    lines += [
        "# Tài Liệu Ôn Tập Module Proposal/Voting",
        "",
        "Dự án: Manga Creation & Publishing Management System",
        "",
        "Mục tiêu của tài liệu này là giúp sinh viên hiểu module `Proposal/Voting` ở mức có thể giải thích với giảng viên, trả lời vấn đáp, phân tích yêu cầu, đưa ra solution, và mô tả được luồng xử lý từ giao diện đến database.",
        "",
        "Tài liệu bám theo source hiện tại trong workspace, gồm các class, method, endpoint, JSP và table thật.",
        "",
        "## 1. Bối Cảnh Nghiệp Vụ",
        "",
        "Module `Proposal` là nơi `MANGAKA` gửi ý tưởng truyện mới cho nhà xuất bản. Một proposal chứa tên truyện, thể loại, tóm tắt nội dung, số chương dự kiến và file mẫu. Hệ thống không cho truyện đi thẳng vào sản xuất, vì ý tưởng phải được kiểm tra trước bởi biên tập viên phụ trách và hội đồng biên tập.",
        "",
        "Bài toán nghiệp vụ mà module giải quyết là kiểm soát chất lượng đầu vào trước khi tạo `Series`. Nếu không có module này, một `MANGAKA` có thể tạo series tùy ý, dẫn tới thiếu kiểm duyệt, không có lịch sử quyết định, không biết ai chịu trách nhiệm, và khó giải thích vì sao truyện được duyệt hoặc bị từ chối.",
        "",
        "Các vai trò chính:",
        "",
        "- `MANGAKA`: người sáng tác, tạo draft, chỉnh sửa proposal khi còn được phép, nộp proposal để review, nhận yêu cầu sửa hoặc kết quả duyệt.",
        "- `TANTOU_EDITOR`: biên tập viên phụ trách, được hệ thống tự động gán theo tải công việc, review bước đầu, quyết định chuyển lên hội đồng, yêu cầu sửa, hoặc từ chối.",
        "- `EDITORIAL_BOARD`: thành viên hội đồng, bỏ phiếu `APPROVE`, `REVISE`, hoặc `REJECT` trong vòng vote.",
        "- `ADMIN`: quản trị hệ thống, xem proposal, cấu hình số lần submit tối đa và số phiếu tối thiểu.",
        "",
        "Vòng đời tổng quát:",
        "",
        "```text",
        "DRAFT",
        "  -> SUBMIT",
        "UNDER_REVIEW",
        "  -> Tantou APPROVE",
        "BOARD_REVIEW",
        "  -> Board vote APPROVE majority",
        "APPROVED -> Series created",
        "",
        "UNDER_REVIEW",
        "  -> Tantou REVISE",
        "REVISION_REQUESTED -> Mangaka edit -> RESUBMIT -> UNDER_REVIEW",
        "",
        "UNDER_REVIEW",
        "  -> Tantou REJECT",
        "REJECTED",
        "",
        "BOARD_REVIEW",
        "  -> Board vote REVISE majority",
        "REVISION_REQUESTED -> Mangaka edit -> RESUBMIT",
        "",
        "BOARD_REVIEW",
        "  -> Board vote REJECT majority",
        "REJECTED",
        "",
        "BOARD_REVIEW",
        "  -> thiếu quorum hoặc hòa",
        "BOARD_REVIEW -> mở vòng vote mới",
        "```",
        "",
        "## 2. Kiến Trúc Tổng Thể",
        "",
        "Luồng chính của module là `JSP view -> Web Controller -> Service -> Repository -> SQL Server`.",
        "",
        "Tầng controller nhận request, đọc `HttpSession`, nhận `@RequestParam`, `@PathVariable`, upload file, nạp model và trả JSP. Tầng service chứa luật nghiệp vụ như role nào được làm gì, quyết định nào hợp lệ, lúc nào cần ghi chú, số lần submit tối đa, và quyền vote. Tầng repository dùng JDBC để đọc/ghi SQL Server, mở transaction, lock proposal khi cần, ghi lịch sử và notification.",
        "",
        "Không để controller truy cập database trực tiếp vì controller chỉ nên lo giao tiếp HTTP. Nếu nhét SQL vào controller, validation và transaction sẽ bị rải rác, API và web khó dùng chung logic, và khi đổi luật nghiệp vụ sẽ phải sửa nhiều nơi.",
        "",
        "## 3. Các File Liên Quan",
        "",
    ]
    lines += table(
        ["Nhóm", "File", "Vai trò"],
        [
            ["Controller web", "`src/java/manga/controller/web/MainController.java`", "Route chính dưới `/main/proposals`, gồm list, create, detail, submit, review, board vote, undo, file download."],
            ["Controller web", "`src/java/manga/controller/web/ProposalController.java`", "Controller Proposal riêng, mapping `/proposals...`; logic gần giống phần Proposal trong `MainController`."],
            ["Controller web", "`src/java/manga/controller/web/ModuleWebController.java`", "Route edit proposal `/main/proposals/{id}/edit` và settings `/main/settings/proposals`."],
            ["Controller API", "`src/java/manga/controller/api/ProposalApiController.java`", "API REST-style dưới `/api/v1/proposals`."],
            ["Service", "`src/java/manga/service/ProposalService.java`", "Luật nghiệp vụ Proposal/Voting."],
            ["Service", "`src/java/manga/service/ProposalSettingsService.java`", "Đọc/ghi cấu hình `maxSubmitAttempts` và `minimumVoteQuorum`."],
            ["Repository", "`src/java/manga/repository/ProposalRepository.java`", "JDBC cho Proposal, vote round, history, notification, tạo Series."],
            ["Repository", "`src/java/manga/repository/SystemSettingRepository.java`", "JDBC cho bảng `SystemSetting` và constraint submit attempts."],
            ["Model", "`src/java/manga/model/Proposal.java`", "Model proposal dùng cho view/API."],
            ["Model", "`src/java/manga/model/ProposalHistory.java`", "Model lịch sử thao tác."],
            ["Model", "`src/java/manga/model/BoardVoteUndoInfo.java`", "Thông tin cửa sổ undo vote 60 giây."],
            ["Model", "`src/java/manga/model/ProposalSettings.java`", "DTO cấu hình proposal."],
            ["Enum", "`src/java/manga/enums/ProposalStatus.java`", "Liệt kê status: `DRAFT`, `UNDER_REVIEW`, `BOARD_REVIEW`, `REVISION_REQUESTED`, `APPROVED`, `REJECTED`."],
            ["Scheduler", "`src/java/manga/scheduler/ProposalBoardVotingScheduler.java`", "Đóng vòng vote hết hạn, nhắc vote, đánh dấu Tantou review quá hạn."],
            ["JSP", "`web/WEB-INF/jsp/proposal/list.jsp`", "Danh sách proposal và trạng thái vote."],
            ["JSP", "`web/WEB-INF/jsp/proposal/detail.jsp`", "Chi tiết proposal, form review, form board vote, history."],
            ["JSP", "`web/WEB-INF/jsp/proposal/create.jsp`", "Form tạo draft."],
            ["JSP", "`web/WEB-INF/jsp/proposal/edit.jsp`", "Form sửa draft hoặc revision."],
            ["JSP", "`web/WEB-INF/jsp/settings/proposals.jsp`", "Trang admin cấu hình Proposal."],
            ["JS/CSS", "`web/assets/js/proposal.js`, `web/assets/css/proposal.css`", "Bắt buộc note theo decision và style vote UI."],
            ["Security", "`src/java/manga/web/interceptor/AuthInterceptor.java`", "RBAC mức URL cho `/main/proposals` và `/main/settings`."],
            ["Database", "`database/schema.sql`, `database/seed_v5.sql`", "Schema và seed data liên quan Proposal/Voting."],
        ],
    )
    lines += [
        "",
        "## 4. Phân Tích Database",
        "",
        "### 4.1 Bảng `Proposal`",
        "",
        "Bảng `Proposal` là trung tâm của module. Mỗi dòng là một ý tưởng truyện do `MANGAKA` tạo.",
        "",
    ]
    lines += table(
        ["Cột", "Ý nghĩa", "Khi nào thay đổi", "Ai dùng"],
        [
            ["`id`", "Khóa chính, định danh proposal.", "Sinh tự động khi `createDraft()` insert.", "Tất cả controller/service/repository."],
            ["`mangakaId`", "Tác giả sở hữu proposal.", "Gán khi tạo draft.", "`getDetail()`, `findForMangaka()`, tạo `Series`."],
            ["`title`", "Tên truyện đề xuất.", "Tạo draft, update draft. Với draft lần đầu có logic khóa identity field ở service/view.", "List/detail/create/edit/Series."],
            ["`genre`", "Thể loại truyện.", "Tạo draft, update draft.", "View, lọc nhận diện nội dung, tạo `Series`."],
            ["`synopsis`", "Tóm tắt nội dung.", "Tạo hoặc sửa.", "Tantou và Board đọc để quyết định."],
            ["`sampleFilePath`", "Đường dẫn file mẫu trong `/uploads/proposals`.", "Tạo hoặc update khi upload file mới.", "Download file, kiểm tra trước submit."],
            ["`originalFileName`", "Tên file gốc từ client.", "Tạo hoặc update khi upload.", "Header download và hiển thị view."],
            ["`approximateChapter`", "Số chương dự kiến.", "Tạo hoặc update.", "Đánh giá quy mô truyện."],
            ["`status`", "Trạng thái workflow.", "Submit, review, vote resolution.", "Tất cả quyền và UI phụ thuộc vào cột này."],
            ["`submittedAt`", "Thời điểm gửi cho Tantou review.", "Set trong `submitForTantouReview()`.", "SLA 48 giờ."],
            ["`rejectedAt`", "Thời điểm bị reject.", "Set khi Tantou hoặc Board reject.", "Theo dõi kết thúc thất bại."],
            ["`assignedEditorId`", "Tantou được gán.", "Set khi submit, chọn bởi `findLeastAssignedTantouEditor()`.", "Review quyền và conflict vote."],
            ["`submitAttemptCount`", "Số lần submit/resubmit.", "Tăng trong `submitForTantouReview()`.", "Giới hạn bởi settings và history."],
            ["`tantouReviewOverdue`", "Cờ quá hạn review 48 giờ.", "Set bởi `markOverdueTantouReviews()`; reset khi resubmit.", "Dashboard/thông báo SLA."],
            ["`createdAt`", "Thời điểm tạo.", "Default `GETDATE()`.", "Sắp xếp danh sách."],
            ["`updatedAt`", "Thời điểm cập nhật cuối.", "Update ở các thao tác ghi.", "Audit nhẹ."],
        ],
    )
    lines += [
        "",
        "Khóa và constraint: `PK_Proposal` trên `id`; `FK_Proposal_Mangaka` trỏ tới `[User](id)`; `FK_Proposal_Editor` trỏ tới `[User](id)`; `CK_Proposal_status` giới hạn status; `CK_Proposal_approxChapter` yêu cầu số chương từ 1; `CK_Proposal_submitAttempts` giới hạn số lần submit.",
        "",
        "### 4.2 Bảng `ProposalHistory`",
        "",
        "`ProposalHistory` là audit log nghiệp vụ của proposal. Điểm đặc biệt: vote của hội đồng không nằm trong bảng `ProposalVote`, mà được ghi vào `ProposalHistory` bằng `actorRole = 'EDITORIAL_BOARD'`, `actionType IN ('APPROVED','REVISE_REQUESTED','REJECTED')` và `boardRoundId`.",
        "",
    ]
    lines += table(
        ["Cột", "Ý nghĩa"],
        [
            ["`id`", "Khóa chính của dòng lịch sử."],
            ["`proposalId`", "Proposal được tác động."],
            ["`actorId`", "User thực hiện; có thể null nếu là `SYSTEM`."],
            ["`actorRole`", "Role tại thời điểm ghi lịch sử: `MANGAKA`, `TANTOU_EDITOR`, `EDITORIAL_BOARD`, `SYSTEM`."],
            ["`actionType`", "Hành động: `CREATED`, `UPDATED`, `SUBMITTED`, `RESUBMITTED`, `ASSIGNED_EDITOR`, `APPROVED`, `REVISE_REQUESTED`, `REJECTED`."],
            ["`note`", "Ghi chú review/vote/reason/system note."],
            ["`submitAttemptNumber`", "Lần submit mà lịch sử này thuộc về."],
            ["`boardRoundId`", "Vòng vote liên quan; null nếu không phải vote board hoặc system event không gắn vòng."],
            ["`createdAt`", "Thời điểm ghi lịch sử."],
        ],
    )
    lines += [
        "",
        "### 4.3 Bảng `ProposalBoardRound`",
        "",
        "Bảng này đại diện cho từng vòng vote của hội đồng. Một proposal có thể có nhiều vòng vote vì thiếu quorum hoặc kết quả hòa.",
        "",
    ]
    lines += table(
        ["Cột", "Ý nghĩa"],
        [
            ["`id`", "Khóa chính vòng vote."],
            ["`proposalId`", "Proposal đang được vote."],
            ["`submitAttemptNumber`", "Lần submit tương ứng."],
            ["`roundNumber`", "Số thứ tự vòng trong cùng proposal/attempt."],
            ["`status`", "`OPEN` hoặc `CLOSED`."],
            ["`openedAt`", "Thời điểm mở vòng."],
            ["`closesAt`", "Hạn đóng, hiện là `DATEADD(DAY, 3, GETDATE())`."],
            ["`closedAt`", "Thời điểm đóng."],
            ["`closeReason`", "Lý do đóng: `EXPIRED`, `ALL_VOTED`."],
        ],
    )
    lines += [
        "",
        "### 4.4 Bảng `ProposalBoardRoundVoter`",
        "",
        "Bảng này chốt danh sách thành viên hội đồng hợp lệ tại thời điểm mở vòng vote. Khóa chính ghép là `(roundId, voterId)`. Thiết kế này giúp vòng vote ổn định ngay cả khi role user thay đổi sau đó.",
        "",
        "### 4.5 Bảng `Series`",
        "",
        "Khi Board approve proposal, `ProposalRepository.ensureSeriesExistsOnApprove()` tạo một dòng trong `Series` nếu chưa có. `Series.proposalId` nối ngược về proposal đã được duyệt, giúp trace series bắt nguồn từ ý tưởng nào.",
        "",
        "### 4.6 Bảng `Notification`",
        "",
        "Module Proposal dùng `Notification` để báo cho Mangaka, Tantou, Board: proposal được mở board review, bị reject, cần revision, vote sắp đóng, review quá hạn, hoặc series đã được tạo.",
        "",
        "### 4.7 Bảng `SystemSetting`",
        "",
        "Hai key quan trọng là `proposal.maxSubmitAttempts` và `proposal.minimumVoteQuorum`. `ProposalSettingsService` đọc settings qua `SystemSettingRepository`; nếu thiếu bảng hoặc thiếu key thì dùng default `2` và `3`.",
        "",
        "Quan hệ chính:",
        "",
        "```text",
        "[User] 1 -- N Proposal.mangakaId",
        "[User] 1 -- N Proposal.assignedEditorId",
        "Proposal 1 -- N ProposalHistory",
        "Proposal 1 -- N ProposalBoardRound",
        "ProposalBoardRound 1 -- N ProposalBoardRoundVoter",
        "ProposalBoardRound 1 -- N ProposalHistory.boardRoundId",
        "Proposal 1 -- 0..1 Series",
        "[User] 1 -- N Notification",
        "```",
        "",
        "## 5. Phân Tích Model `Proposal`",
        "",
        "`Proposal.java` không chỉ map cột bảng `Proposal`; nó còn chứa field tổng hợp cho UI vote như `boardApproveVotes`, `boardRoundStatus`, `boardEligibleVoterCount`. Các field này được tính bằng subquery trong `ProposalRepository.BOARD_SELECT_COLUMNS`.",
        "",
    ]
    lines += table(
        ["Field", "Nguồn dữ liệu", "Giải thích nhanh"],
        [
            ["`id`", "`Proposal.id`", "Định danh proposal."],
            ["`mangakaId`", "`Proposal.mangakaId`", "Người sở hữu."],
            ["`title`, `genre`, `synopsis`", "`Proposal`", "Nội dung chính để review."],
            ["`sampleFilePath`, `originalFileName`", "`Proposal`", "Lưu và tải file mẫu."],
            ["`approximateChapter`", "`Proposal.approximateChapter`", "Quy mô dự kiến."],
            ["`status`", "`Proposal.status`", "Điều khiển workflow và quyền."],
            ["`submittedAt`, `rejectedAt`", "`Proposal`", "Mốc thời gian nghiệp vụ."],
            ["`assignedEditorId`", "`Proposal.assignedEditorId`", "Tantou đang quản lý."],
            ["`submitAttemptCount`", "`Proposal.submitAttemptCount`", "Giới hạn số lần resubmit."],
            ["`boardApproveVotes`, `boardReviseVotes`, `boardRejectVotes`, `boardTotalVotes`", "Subquery `ProposalHistory`", "Số phiếu theo decision trong vòng hiện tại."],
            ["`boardRoundId`, `boardRoundNumber`, `boardRoundStatus`", "Subquery `ProposalBoardRound`", "Vòng vote hiện tại hoặc mới nhất."],
            ["`boardVotingOpenedAt`, `boardVotingClosesAt`, `boardVotingClosedAt`", "Subquery `ProposalBoardRound`", "Thông tin thời gian vote."],
            ["`boardEligibleVoterCount`", "Subquery `ProposalBoardRoundVoter`", "Số voter của vòng."],
        ],
    )
    lines += [
        "",
        "## 6. Controller Và Endpoint",
        "",
        "### 6.1 `MainController`",
        "",
        "`MainController` có `@RequestMapping('/main')`, nên các endpoint thật trên web thường bắt đầu bằng `/main`.",
        "",
    ]
    lines += table(
        ["Endpoint", "Method", "Method Java", "Input", "Kết quả"],
        [
            ["`/main/proposals`", "GET", "`proposals()`", "`HttpSession`, `Model`", "Load list qua `proposalService.listForUser()`, trả `proposal/list`."],
            ["`/main/proposals/create`", "GET", "`createProposalPage()`", "session user", "Chỉ `MANGAKA`; trả `proposal/create`."],
            ["`/main/proposals/create`", "POST", "`createProposal()`", "`title`, `genre`, `synopsis`, `approximateChapter`, file `sampleFile`", "Lưu file rồi gọi `createProposal()`, redirect detail."],
            ["`/main/proposals/{id}`", "GET", "`proposalDetail()`", "`id`, session", "Load detail, history, voter list, quyền edit/review/vote, trả `proposal/detail`."],
            ["`/main/proposals/{id}/vote`", "GET", "`proposalVoteDeepLink()`", "`id`", "Deep link về detail để Board vote."],
            ["`/main/proposals/{id}/file`", "GET", "`proposalFile()`", "`id`", "Kiểm quyền qua `getDetail()`, stream file upload."],
            ["`/main/proposals/{id}/submit`", "POST", "`submitProposal()`", "`id`", "Mangaka submit draft/revision."],
            ["`/main/proposals/{id}/review`", "POST", "`reviewProposal()`", "`decision`, `note`", "Tantou review."],
            ["`/main/proposals/{id}/board-vote`", "POST", "`boardVoteProposal()`", "`decision`, `note`", "Board vote."],
            ["`/main/proposals/{id}/board-vote/undo`", "POST", "`undoBoardVote()`", "`id`", "Undo vote trong 60 giây."],
        ],
    )
    lines += [
        "",
        "### 6.2 `ModuleWebController`",
        "",
    ]
    lines += table(
        ["Endpoint", "Method Java", "Vai trò"],
        [
            ["`GET /main/proposals/{id}/edit`", "`proposalEditPage()`", "Mở form sửa nếu user là chủ proposal, role `MANGAKA`, status editable, chưa quá giới hạn submit."],
            ["`POST /main/proposals/{id}/edit`", "`proposalEdit()`", "Lưu sửa proposal qua `proposalService.updateDraft()`."],
            ["`GET /main/settings/proposals`", "`proposalSettings()`", "Admin xem cấu hình."],
            ["`POST /main/settings/proposals`", "`proposalSettingsSave()`", "Admin cập nhật `maxSubmitAttempts` và `minimumVoteQuorum`."],
        ],
    )
    lines += [
        "",
        "### 6.3 `ProposalController`",
        "",
        "`ProposalController` mapping `/proposals...` không có prefix `/main`. Logic gần giống `MainController`: `list()`, `createPage()`, `create()`, `detail()`, `submit()`, `review()`, `boardVote()`, `undoBoardVote()`, `downloadFile()`. Trong source hiện tại, giao diện header và notification chủ yếu trỏ về `/main/proposals`, nên khi bảo vệ cần nói route chính là `MainController`, còn `ProposalController` là controller Proposal riêng có mapping không prefix.",
        "",
        "### 6.4 `ProposalApiController`",
        "",
    ]
    lines += table(
        ["Endpoint", "Method", "Method Java", "Ý nghĩa"],
        [
            ["`/api/v1/proposals`", "GET", "`list()`", "Trả danh sách proposal theo role."],
            ["`/api/v1/proposals`", "POST", "`create()`", "Tạo draft bằng API, nhận `sampleFilePath` thay vì multipart upload."],
            ["`/api/v1/proposals/{id}`", "GET", "`detail()`", "Trả chi tiết."],
            ["`/api/v1/proposals/{id}`", "PUT", "`updateDraft()`", "Cập nhật draft/revision."],
            ["`/api/v1/proposals/{id}/submit`", "POST", "`submit()`", "Submit proposal."],
            ["`/api/v1/proposals/{id}/review`", "POST", "`review()`", "Tantou review."],
            ["`/api/v1/proposals/{id}/board-vote`", "POST", "`boardVote()`", "Board vote và trả `BoardVoteUndoInfo`."],
            ["`/api/v1/proposals/{id}/board-vote/undo`", "PATCH", "`undoBoardVote()`", "Undo vote."],
            ["`/api/v1/proposals/{id}/my-board-vote`", "GET", "`myBoardVote()`", "Xem cửa sổ undo vote của user hiện tại."],
        ],
    )
    lines += [
        "",
        "## 7. Phân Tích `ProposalService`",
        "",
    ]
    lines += table(
        ["Method", "Mục đích", "Luật nghiệp vụ chính"],
        [
            ["`listGenres()`", "Trả danh sách genre cố định.", "Genre hợp lệ phải nằm trong list này."],
            ["`listForUser(user)`", "Danh sách proposal theo role.", "`MANGAKA` chỉ thấy của mình; `EDITORIAL_BOARD`/`ADMIN` thấy proposal không draft; `TANTOU_EDITOR` thấy proposal được assign."],
            ["`getDetail(user, proposalId)`", "Load chi tiết và check quyền.", "Mangaka chỉ xem proposal của mình; Board/Admin không xem draft; Tantou chỉ xem proposal được assign."],
            ["`createProposal(...)`", "Tạo draft.", "Chỉ `MANGAKA`; nội dung hợp lệ; không có active draft khác."],
            ["`updateDraft(...)`", "Cập nhật draft hoặc revision.", "Chỉ owner `MANGAKA`; nếu draft lần đầu thì title/genre bị giữ nguyên theo dữ liệu cũ; file có thể giữ file cũ."],
            ["`submitProposal(...)`", "Submit sang Tantou review.", "Chỉ `MANGAKA`; không có active proposal khác; chưa quá max attempts; phải có file và chapter."],
            ["`reviewProposal(...)`", "Tantou quyết định.", "Chỉ `TANTOU_EDITOR`; decision phải là `APPROVE`, `REJECT`, `REVISE`; reject/revise bắt buộc note."],
            ["`voteProposalAsBoard(...)`", "Board bỏ phiếu.", "Chỉ `EDITORIAL_BOARD`; proposal phải `BOARD_REVIEW`; không được vote proposal mình đang quản lý với vai trò Tantou; không vote trùng; revise/reject cần note."],
            ["`undoBoardVote(...)`", "Rút phiếu.", "Chỉ `EDITORIAL_BOARD`; repository cho undo trong 60 giây khi proposal vẫn `BOARD_REVIEW`."],
            ["`canCastBoardVote(user, proposal)`", "Tính quyền hiển thị form vote.", "Kiểm role, status, roundId, conflict Tantou, chưa vote."],
            ["`boardVoteBlockMessage(user, proposal)`", "Thông báo lý do không vote.", "Hiện dùng cho conflict: Tantou không được vote proposal mình quản lý."],
            ["`getBoardVoteUndoInfo(user, proposalId)`", "Trả thông tin undo.", "Chỉ khi vote mới nhất còn trong 60 giây."],
            ["`listBoardRoundVoters(user, proposalId)`", "Danh sách voter của round.", "Gọi `getDetail()` trước để check quyền."],
            ["`listHistory(user, proposalId)`", "Lịch sử proposal.", "Role hợp lệ mới xem được."],
            ["`getMaxSubmitAttempts()`", "Đọc setting submit attempts.", "Default 2 nếu chưa có setting."],
            ["`getMinimumVoteQuorum()`", "Đọc setting quorum.", "Default 3 nếu chưa có setting."],
        ],
    )
    lines += [
        "",
        "Quy tắc quan trọng: `TANTOU_EDITOR` không được vote proposal do chính mình quản lý. Lý do là tránh xung đột lợi ích: người đã review bước đầu có thể thiên vị khi vote hội đồng. Code kiểm tra trong `ProposalService.isManagingTantouForProposal()`, `boardVoteBlockMessage()`, `canCastBoardVote()` và `assertCanCastBoardVote()`.",
        "",
        "## 8. Phân Tích `ProposalRepository`",
        "",
        "`ProposalRepository` là nơi có phần lớn logic transaction. Các method ghi dữ liệu thường dùng `conn.setAutoCommit(false)`, `commit()` khi thành công và `rollback()` khi lỗi. Khi thay đổi trạng thái quan trọng, repository dùng `lockProposal()` với `WITH (UPDLOCK, ROWLOCK)` để tránh hai request xử lý cùng proposal một lúc.",
        "",
    ]
    lines += table(
        ["Method", "Tác động database", "Giải thích"],
        [
            ["`findForMangaka()`", "SELECT `Proposal` theo `mangakaId`.", "Trang list của Mangaka."],
            ["`findForBoardAndEditor()`", "SELECT proposal status không phải draft.", "Admin/Board xem pipeline review."],
            ["`findForAssignedEditor()`", "SELECT proposal được assign cho editor.", "Tantou chỉ thấy việc của mình."],
            ["`findById()`", "SELECT 1 proposal.", "Dùng cho detail và validation."],
            ["`createDraft()`", "INSERT `Proposal`; INSERT `ProposalHistory` action `CREATED`.", "Tạo draft trong transaction."],
            ["`updateDraft()`", "UPDATE `Proposal`; INSERT history `UPDATED`.", "Chỉ update owner và status `DRAFT`/`REVISION_REQUESTED`."],
            ["`submitForTantouReview()`", "Lock proposal; UPDATE status `UNDER_REVIEW`; assign Tantou; tăng attempt; INSERT history.", "Tự động gán Tantou ít việc nhất."],
            ["`reviewByTantou()`", "Update status theo decision; mở board round nếu approve; notification.", "Bước review đầu tiên."],
            ["`hasBoardVote()`", "COUNT vote trong current round.", "Chống vote trùng."],
            ["`findLatestBoardVote()`", "SELECT vote mới nhất của actor.", "Dùng cho undo window."],
            ["`undoBoardVote()`", "DELETE dòng `ProposalHistory` vote.", "Chỉ cho undo trong 60 giây."],
            ["`voteByEditorialBoard()`", "Lock proposal; insert vote history; đóng round nếu all voted.", "Board vote có kiểm round mở, chưa hết hạn, eligible voter."],
            ["`closeExpiredBoardVotingRounds()`", "Tìm round hết hạn; đóng và resolve.", "Scheduler gọi 15 phút/lần."],
            ["`notifyBoardVotingClosingSoon()`", "INSERT notification cho voter chưa vote.", "Scheduler gọi hourly khi còn <=24 giờ."],
            ["`markOverdueTantouReviews()`", "UPDATE `tantouReviewOverdue`; INSERT notification.", "SLA review 48 giờ."],
            ["`listBoardRoundVoters()`", "SELECT voter, vote status, conflict.", "Hiển thị danh sách Board member trong detail."],
            ["`listHistory()`", "SELECT history join `[User]`.", "Bảng lịch sử trên detail."],
            ["`hasActiveProposal()`", "COUNT proposal `UNDER_REVIEW`, `BOARD_REVIEW`, `REVISION_REQUESTED`.", "Không cho Mangaka có nhiều proposal active."],
            ["`hasActiveDraft()`", "COUNT proposal `DRAFT`.", "Không cho nhiều draft active."],
            ["`canVoteInCurrentBoardRound()`", "COUNT voter đủ điều kiện và chưa vote.", "Hàm tiện ích kiểm eligible."],
            ["`ensureSeriesExistsOnApprove()`", "INSERT `Series` nếu chưa có.", "Board approve thì tạo series."],
            ["`openBoardVotingRound()`", "INSERT `ProposalBoardRound`; INSERT `ProposalBoardRoundVoter`; INSERT system history.", "Mở vòng vote 3 ngày."],
            ["`resolveClosedBoardRound()`", "COUNT vote; update status; tạo series hoặc mở round mới.", "Luật kết quả vote."],
        ],
    )
    lines += [
        "",
        "Luật resolve board vote trong `resolveClosedBoardRound()`: nếu tổng phiếu nhỏ hơn `minimumVoteQuorum` thì mở vòng mới; nếu mỗi loại approve/revise/reject đều có 1 phiếu thì coi là hòa đặc biệt và mở vòng mới; nếu một loại phiếu lớn hơn hai loại còn lại thì quyết định theo loại đó; các trường hợp hòa khác cũng mở vòng mới.",
        "",
        "## 9. Frontend JSP",
        "",
    ]
    lines += table(
        ["JSP", "Hiển thị", "Form/Button quan trọng"],
        [
            ["`proposal/list.jsp`", "Bảng proposal: title, genre, approximate chapter, status, board voting progress.", "`New Proposal` tới `/main/proposals/create`; `View` tới `/main/proposals/{id}`."],
            ["`proposal/create.jsp`", "Form tạo draft.", "POST `/main/proposals/create`, enctype multipart, field `sampleFile` bắt buộc."],
            ["`proposal/edit.jsp`", "Form sửa draft/revision.", "POST `/main/proposals/{id}/edit`; file optional; title/genre có thể readonly/disabled khi `lockIdentityFields`."],
            ["`proposal/detail.jsp`", "Thông tin proposal, file, synopsis, voting summary, voter list, history.", "Submit to Tantou, Tantou Review, Board Vote, Undo vote nếu có."],
            ["`settings/proposals.jsp`", "Form admin chỉnh settings.", "POST `/main/settings/proposals`."],
            ["`dashboard/index.jsp`", "Metric proposal và active proposal.", "Link View Details tới `/main/proposals/{id}`."],
            ["`common/header.jsp`", "Navigation.", "Hiện menu Proposals cho Admin/Mangaka/Tantou/Board; Settings cho Admin."],
        ],
    )
    lines += [
        "",
        "`web/assets/js/proposal.js` làm hai việc nhỏ nhưng quan trọng: note Tantou bắt buộc khi decision là `REJECT` hoặc `REVISE`; note Board bắt buộc khi decision là `REVISE` hoặc `REJECT`; nút lý do reject hiển thị alert từ `data-reject-reason`.",
        "",
        "## 10. Request Lifecycle",
        "",
        "### 10.1 Mangaka tạo proposal",
        "",
        "Browser mở `GET /main/proposals/create`; `MainController.createProposalPage()` kiểm user có role `MANGAKA`, đưa `genres` vào model và trả `proposal/create.jsp`. Khi submit form, browser gửi `POST /main/proposals/create` kèm multipart file. Controller gọi `saveUpload()` để lưu file vào `/uploads/proposals`, sanitize tên file, giới hạn 20 MB, rồi gọi `ProposalService.createProposal()`. Service kiểm role, kiểm content, kiểm active draft. Repository `createDraft()` insert `Proposal` status `DRAFT`, insert history `CREATED`, commit transaction. Controller redirect về `/main/proposals/{id}`.",
        "",
        "### 10.2 Mangaka submit proposal",
        "",
        "Từ detail, form gửi `POST /main/proposals/{id}/submit`. Service `submitProposal()` kiểm role `MANGAKA`, proposal tồn tại, không có active proposal khác, chưa quá `maxSubmitAttempts`, có file và approximate chapter. Repository `submitForTantouReview()` lock proposal, kiểm owner và status editable, chọn Tantou bằng `findLeastAssignedTantouEditor()`, update status `UNDER_REVIEW`, set `submittedAt`, set `assignedEditorId`, tăng `submitAttemptCount`, reset `tantouReviewOverdue`, ghi history `SUBMITTED` hoặc `RESUBMITTED`, ghi system history `ASSIGNED_EDITOR`.",
        "",
        "### 10.3 Tantou review proposal",
        "",
        "Tantou mở detail. `MainController.proposalDetail()` đặt `canReview = true` nếu user có role `TANTOU_EDITOR`, là editor được assign, và status `UNDER_REVIEW`. Form gửi `POST /main/proposals/{id}/review`. Service chuẩn hóa decision, chỉ nhận `APPROVE`, `REJECT`, `REVISE`; reject/revise phải có note. Repository `reviewByTantou()` lock proposal, kiểm assigned editor, kiểm status. Nếu approve: status sang `BOARD_REVIEW`, ghi history `APPROVED`, gọi `openBoardVotingRound()`, gửi notification cho Board. Nếu reject: status `REJECTED`, set `rejectedAt`, ghi history, notify Mangaka. Nếu revise: status `REVISION_REQUESTED`, ghi history, notify Mangaka.",
        "",
        "### 10.4 Board vote proposal",
        "",
        "Board mở detail. Service `canCastBoardVote()` kiểm user có role `EDITORIAL_BOARD`, proposal status `BOARD_REVIEW`, có `boardRoundId`, không phải Tantou đang quản lý proposal đó, và chưa vote. Form gửi `POST /main/proposals/{id}/board-vote`. Service kiểm decision và note. Repository `voteByEditorialBoard()` lock proposal, tìm open round, kiểm round chưa hết hạn, kiểm voter có trong `ProposalBoardRoundVoter`, kiểm chưa vote, insert vote vào `ProposalHistory`. Nếu tất cả eligible board members đã vote, repository đóng round với `ALL_VOTED` và resolve ngay.",
        "",
        "### 10.5 Proposal approved và tạo series",
        "",
        "Khi round đóng, `resolveClosedBoardRound()` đếm phiếu trong `ProposalHistory`. Nếu `APPROVED` có số phiếu lớn nhất và đủ quorum, repository gọi `ensureSeriesExistsOnApprove()`. Hàm này kiểm `Series` đã tồn tại chưa; nếu chưa thì insert `Series` với `proposalId`, `mangakaId`, `tantouEditorId`, `title`, `genre`, status `ACTIVE`. Sau đó update proposal status `APPROVED`, ghi history system `APPROVED`, gửi notification cho Mangaka và Tantou.",
        "",
        "## 11. Phân Quyền",
        "",
    ]
    lines += table(
        ["Role", "Được phép", "Bị cấm/giới hạn", "Nơi check"],
        [
            ["`MANGAKA`", "Tạo, sửa, submit proposal của mình; xem history của mình.", "Không xem proposal người khác; không review/vote.", "`AuthInterceptor`, `ProposalService`, JSP `canEdit`, `canSubmit`."],
            ["`TANTOU_EDITOR`", "Xem proposal được assign; review proposal `UNDER_REVIEW`.", "Không review proposal chưa assign; không vote proposal mình quản lý nếu đồng thời là Board.", "`getDetail()`, `reviewByTantou()`, `canCastBoardVote()`."],
            ["`EDITORIAL_BOARD`", "Xem proposal không draft; vote proposal trong board round.", "Không xem draft; không vote trùng; không vote nếu conflict Tantou.", "`getDetail()`, `voteProposalAsBoard()`, `ProposalBoardRoundVoter`."],
            ["`ADMIN`", "Xem proposal không draft; chỉnh settings.", "Không có flow tạo/review/vote nghiệp vụ trong service trừ quyền xem.", "`AuthInterceptor`, `ModuleWebController.requireAdmin()`."],
        ],
    )
    lines += [
        "",
        "`AuthInterceptor.isAllowed()` chặn URL trước controller: `/main/proposals` cho Admin/Mangaka/Tantou/Board; `/main/settings` chỉ Admin; API `/api/v1/users` chỉ Admin. Sau đó service vẫn kiểm quyền chi tiết theo dữ liệu từng proposal.",
        "",
        "## 12. Status Transition",
        "",
    ]
    lines += table(
        ["Từ", "Sang", "Ai kích hoạt", "Method xử lý", "Ghi chú"],
        [
            ["-", "`DRAFT`", "`MANGAKA`", "`createProposal()` -> `createDraft()`", "Tạo proposal mới."],
            ["`DRAFT`", "`UNDER_REVIEW`", "`MANGAKA`", "`submitProposal()` -> `submitForTantouReview()`", "Lần submit đầu."],
            ["`REVISION_REQUESTED`", "`UNDER_REVIEW`", "`MANGAKA`", "`submitProposal()` -> `submitForTantouReview()`", "Resubmit sau sửa."],
            ["`UNDER_REVIEW`", "`BOARD_REVIEW`", "`TANTOU_EDITOR`", "`reviewProposal()` -> `reviewByTantou(APPROVE)`", "Mở board round."],
            ["`UNDER_REVIEW`", "`REVISION_REQUESTED`", "`TANTOU_EDITOR`", "`reviewByTantou(REVISE)`", "Yêu cầu sửa."],
            ["`UNDER_REVIEW`", "`REJECTED`", "`TANTOU_EDITOR`", "`reviewByTantou(REJECT)`", "Reject ở tầng Tantou."],
            ["`BOARD_REVIEW`", "`APPROVED`", "System sau Board vote", "`resolveClosedBoardRound()`", "Approve thắng, tạo series."],
            ["`BOARD_REVIEW`", "`REVISION_REQUESTED`", "System sau Board vote", "`resolveClosedBoardRound()`", "Revise thắng."],
            ["`BOARD_REVIEW`", "`REJECTED`", "System sau Board vote", "`resolveClosedBoardRound()`", "Reject thắng."],
            ["`BOARD_REVIEW`", "`BOARD_REVIEW`", "System", "`resolveClosedBoardRound()` -> `openBoardVotingRound()`", "Thiếu quorum hoặc hòa."],
        ],
    )
    lines += [
        "",
        "## 13. Scheduler",
        "",
    ]
    lines += table(
        ["Scheduler method", "Cron", "Repository method", "Mục đích"],
        [
            ["`closeExpiredBoardVotingRounds()`", "`0 */15 * * * *`", "`ProposalRepository.closeExpiredBoardVotingRounds()`", "Mỗi 15 phút đóng các vòng vote đã quá `closesAt` và resolve kết quả."],
            ["`remindBoardVotingClosingSoon()`", "`0 0 * * * *`", "`ProposalRepository.notifyBoardVotingClosingSoon()`", "Mỗi giờ gửi notification cho Board chưa vote khi còn <=24 giờ."],
            ["`markOverdueTantouReviews()`", "`0 5 * * * *`", "`ProposalRepository.markOverdueTantouReviews()`", "Mỗi giờ phút 5 đánh dấu proposal `UNDER_REVIEW` quá 48 giờ và notify Tantou."],
        ],
    )
    lines += [
        "",
        "Scheduler cần thiết vì không phải lúc nào cũng có user request đúng thời điểm deadline. Nếu chỉ xử lý khi user mở trang, round hết hạn có thể treo mãi, SLA quá hạn không được báo, và hệ thống không chủ động nhắc người chưa vote.",
        "",
        "## 14. Mapping Business -> Code",
        "",
    ]
    lines += table(
        ["Yêu cầu nghiệp vụ", "Controller", "Service", "Repository", "Table"],
        [
            ["Mangaka tạo proposal", "`MainController.createProposal()`", "`ProposalService.createProposal()`", "`ProposalRepository.createDraft()`", "`Proposal`, `ProposalHistory`"],
            ["Mangaka sửa draft/revision", "`ModuleWebController.proposalEdit()`", "`ProposalService.updateDraft()`", "`ProposalRepository.updateDraft()`", "`Proposal`, `ProposalHistory`"],
            ["Mangaka submit/resubmit", "`MainController.submitProposal()`", "`ProposalService.submitProposal()`", "`ProposalRepository.submitForTantouReview()`", "`Proposal`, `ProposalHistory`"],
            ["Tự gán Tantou ít việc", "-", "-", "`findLeastAssignedTantouEditor()`", "`User`, `UserRole`, `Role`, `Proposal`"],
            ["Tantou approve lên Board", "`MainController.reviewProposal()`", "`reviewProposal()`", "`reviewByTantou()`, `openBoardVotingRound()`", "`Proposal`, `ProposalBoardRound`, `ProposalBoardRoundVoter`, `ProposalHistory`, `Notification`"],
            ["Tantou reject", "`reviewProposal()`", "`reviewProposal()`", "`reviewByTantou()`", "`Proposal`, `ProposalHistory`, `Notification`"],
            ["Tantou yêu cầu sửa", "`reviewProposal()`", "`reviewProposal()`", "`reviewByTantou()`", "`Proposal`, `ProposalHistory`, `Notification`"],
            ["Board vote", "`boardVoteProposal()`", "`voteProposalAsBoard()`", "`voteByEditorialBoard()`", "`ProposalHistory`, `ProposalBoardRound`"],
            ["Undo vote 60 giây", "`undoBoardVote()`", "`undoBoardVote()`", "`undoBoardVote()`", "`ProposalHistory`"],
            ["Đóng round hết hạn", "`ProposalBoardVotingScheduler`", "-", "`closeExpiredBoardVotingRounds()`", "`ProposalBoardRound`, `ProposalHistory`, `Proposal`, `Series`"],
            ["Tạo series khi approved", "-", "-", "`ensureSeriesExistsOnApprove()`", "`Series`, `Proposal`"],
            ["Admin chỉnh settings", "`ModuleWebController.proposalSettingsSave()`", "`ProposalSettingsService.updateSettings()`", "`SystemSettingRepository.setProposalSettings()`", "`SystemSetting`, `Proposal` constraint"],
        ],
    )
    lines += [
        "",
        "## 15. Review Thiết Kế Hệ Thống",
        "",
        "Điểm mạnh:",
        "",
        "- Có phân tầng rõ: controller, service, repository.",
        "- Có transaction ở repository cho các thao tác nhiều bước.",
        "- Có audit trail bằng `ProposalHistory`.",
        "- Board voting có vòng vote riêng, deadline, quorum, reopen khi thiếu phiếu/hòa.",
        "- Có `AuthInterceptor` kiểm URL-level RBAC và service kiểm dữ liệu chi tiết.",
        "- Có scheduler để xử lý deadline và SLA.",
        "",
        "Điểm yếu:",
        "",
        "- Vote được lưu trong `ProposalHistory` thay vì bảng `ProposalVote` riêng, khiến history vừa là audit log vừa là dữ liệu quyết định. Điều này tiết kiệm bảng nhưng query phức tạp hơn.",
        "- Có logic Proposal bị trùng giữa `MainController` và `ProposalController`, dễ lệch behavior khi sửa một bên.",
        "- `ProposalRepository` chứa nhiều trách nhiệm: proposal CRUD, voting, notification, series creation, settings awareness. Có thể tách thành `ProposalVoteRepository`, `ProposalNotificationRepository`, `SeriesCreationService`.",
        "- Một số thông báo và comment trong source bị lỗi encoding, dễ gây khó đọc.",
        "- Upload file lưu dưới webroot, cần kiểm kỹ loại file và bảo mật download.",
        "",
        "Vấn đề bảo mật cần chú ý:",
        "",
        "- Cần validate file type chứ hiện mới giới hạn size và sanitize tên file.",
        "- `originalFileName` được đưa vào header download; nên escape/normalize kỹ hơn.",
        "- API tạo proposal nhận `sampleFilePath` trực tiếp, có nguy cơ client truyền path không hợp lệ nếu không có tầng upload riêng.",
        "- Cần CSRF protection cho POST web form nếu deploy thật.",
        "",
        "Vấn đề hiệu năng:",
        "",
        "- `BOARD_SELECT_COLUMNS` dùng nhiều subquery cho mỗi proposal; danh sách lớn có thể chậm.",
        "- Nên index `Proposal(status, assignedEditorId)`, `ProposalHistory(boardRoundId, actorRole, actionType)`, `ProposalBoardRound(proposalId, submitAttemptNumber, roundNumber)`.",
        "",
        "Khả năng mở rộng:",
        "",
        "- Có thể thêm bảng `ProposalVote` để lưu vote rõ hơn, giữ `ProposalHistory` chỉ làm audit.",
        "- Có thể thêm policy engine cho status transition.",
        "- Có thể chuyển notification sang service riêng để repository không ôm side effect quá nhiều.",
        "",
        "Design pattern có thể nhận ra:",
        "",
        "- Layered Architecture: controller/service/repository.",
        "- Repository Pattern: `ProposalRepository` đóng gói JDBC.",
        "- DTO/View Model: `Proposal`, `ProposalHistory`, `BoardVoteUndoInfo` được dùng để truyền dữ liệu ra view/API.",
        "- Scheduler/Background Job: `ProposalBoardVotingScheduler` xử lý deadline ngoài request lifecycle.",
        "",
        "## 16. Cheat Sheet Ôn Nhanh 15 Phút",
        "",
        "Table quan trọng: `Proposal`, `ProposalHistory`, `ProposalBoardRound`, `ProposalBoardRoundVoter`, `Series`, `Notification`, `SystemSetting`.",
        "",
        "Endpoint quan trọng: `/main/proposals`, `/main/proposals/create`, `/main/proposals/{id}`, `/main/proposals/{id}/edit`, `/main/proposals/{id}/submit`, `/main/proposals/{id}/review`, `/main/proposals/{id}/board-vote`, `/main/proposals/{id}/board-vote/undo`, `/main/settings/proposals`, `/api/v1/proposals`.",
        "",
        "Status quan trọng: `DRAFT`, `UNDER_REVIEW`, `BOARD_REVIEW`, `REVISION_REQUESTED`, `APPROVED`, `REJECTED`.",
        "",
        "Business rule quan trọng: chỉ `MANGAKA` tạo/submit; chỉ owner được sửa; không quá max submit attempts; submit phải có file và approximate chapter; Tantou chỉ review proposal được assign; reject/revise cần note; Board không vote trùng; Tantou không vote proposal mình quản lý; vote có undo 60 giây; round kéo dài 3 ngày; thiếu quorum hoặc hòa thì mở round mới; approve thắng thì tạo `Series`.",
        "",
        "Scheduler quan trọng: đóng round hết hạn mỗi 15 phút; nhắc vote sắp đóng mỗi giờ; đánh dấu Tantou review quá hạn 48 giờ mỗi giờ.",
        "",
        "Câu dễ bị hỏi: vì sao cần service layer, vì sao cần history, vì sao vote gắn board round, vì sao cần quorum, vì sao Tantou không được vote proposal của mình, vì sao cần scheduler, vì sao tạo series ở repository khi approve.",
        "",
        "## 17. 100 Câu Hỏi Vấn Đáp Có Đáp Án Mẫu",
        "",
    ]
    questions = [
        ("Module `Proposal` dùng để làm gì?", "Module này quản lý quá trình Mangaka gửi ý tưởng truyện mới, được Tantou review, được Editorial Board vote, rồi mới có thể tạo `Series`."),
        ("Vì sao hệ thống không cho Mangaka tạo `Series` trực tiếp?", "Vì cần kiểm duyệt chất lượng, lưu lịch sử quyết định, phân công editor, và đảm bảo hội đồng có quyền quyết định trước khi đưa truyện vào sản xuất."),
        ("Status `DRAFT` nghĩa là gì?", "`DRAFT` là proposal còn ở bản nháp, do Mangaka sở hữu, chưa gửi cho Tantou review."),
        ("Status `UNDER_REVIEW` nghĩa là gì?", "`UNDER_REVIEW` là proposal đã submit và đang chờ Tantou Editor được assign review."),
        ("Status `BOARD_REVIEW` nghĩa là gì?", "`BOARD_REVIEW` là proposal đã qua Tantou và đang trong vòng vote của Editorial Board."),
        ("Status `REVISION_REQUESTED` nghĩa là gì?", "Proposal cần Mangaka chỉnh sửa theo note của Tantou hoặc Board trước khi resubmit."),
        ("Status `APPROVED` nghĩa là gì?", "Proposal đã được chấp thuận để xuất bản và hệ thống đã hoặc sẽ tạo `Series`."),
        ("Status `REJECTED` nghĩa là gì?", "Proposal bị từ chối, không tiếp tục trong workflow hiện tại."),
        ("Class nào chứa business logic chính?", "`ProposalService.java` chứa logic quyền, validation và điều phối sang repository."),
        ("Class nào chứa SQL chính?", "`ProposalRepository.java` chứa hầu hết SQL liên quan proposal, vote round, history, notification và tạo series."),
        ("Vì sao cần `ProposalHistory`?", "Để lưu audit trail: ai làm gì, role nào, note gì, ở attempt nào, lúc nào."),
        ("Board vote được lưu ở đâu?", "Vote được lưu trong `ProposalHistory` với `actorRole = 'EDITORIAL_BOARD'`, `actionType` là `APPROVED`, `REVISE_REQUESTED` hoặc `REJECTED`, và có `boardRoundId`."),
        ("Vì sao cần `ProposalBoardRound`?", "Để quản lý từng vòng vote, deadline 3 ngày, trạng thái mở/đóng và khả năng mở vòng mới khi thiếu quorum hoặc hòa."),
        ("Vì sao cần `ProposalBoardRoundVoter`?", "Để chốt danh sách người được vote trong từng vòng, tránh thay đổi role sau đó làm sai kết quả."),
        ("Ai được tạo proposal?", "Chỉ user có role `MANGAKA`, kiểm trong `ProposalService.createProposal()`."),
        ("Ai được review ở tầng Tantou?", "Chỉ `TANTOU_EDITOR` được assign cho proposal đó, kiểm trong `reviewByTantou()`."),
        ("Ai được vote Board?", "User có role `EDITORIAL_BOARD`, nằm trong `ProposalBoardRoundVoter`, proposal đang `BOARD_REVIEW`, chưa vote và không conflict."),
        ("Vì sao Tantou không được vote proposal mình quản lý?", "Để tránh xung đột lợi ích vì Tantou đã có ảnh hưởng ở bước review đầu."),
        ("Code nào kiểm conflict Tantou khi vote?", "`ProposalService.isManagingTantouForProposal()`, `boardVoteBlockMessage()`, `assertCanCastBoardVote()`."),
        ("Reject hoặc revise có cần note không?", "Có. `ProposalService.reviewProposal()` và `voteProposalAsBoard()` bắt buộc note khi decision là reject/revise."),
        ("Submit proposal cần điều kiện gì?", "User là `MANGAKA`, proposal tồn tại, không có active proposal khác, chưa quá max attempts, có file và approximate chapter."),
        ("Active proposal gồm status nào?", "`UNDER_REVIEW`, `BOARD_REVIEW`, `REVISION_REQUESTED`, kiểm bởi `hasActiveProposal()`."),
        ("Active draft được kiểm thế nào?", "`hasActiveDraft()` đếm proposal status `DRAFT` của Mangaka, loại trừ proposal hiện tại nếu có."),
        ("Tantou được assign bằng cách nào?", "`findLeastAssignedTantouEditor()` chọn active Tantou có ít proposal active nhất."),
        ("Khi Tantou approve thì chuyện gì xảy ra?", "Proposal sang `BOARD_REVIEW`, ghi history `APPROVED`, mở board round, notify Board."),
        ("Khi Tantou reject thì chuyện gì xảy ra?", "Proposal sang `REJECTED`, set `rejectedAt`, ghi history, notify Mangaka."),
        ("Khi Tantou revise thì chuyện gì xảy ra?", "Proposal sang `REVISION_REQUESTED`, ghi history, notify Mangaka để sửa."),
        ("Board round kéo dài bao lâu?", "3 ngày, do `openBoardVotingRound()` set `closesAt = DATEADD(DAY, 3, GETDATE())`."),
        ("Khi nào round đóng ngay?", "Khi tất cả eligible board members đã vote, `allEligibleBoardMembersVoted()` trả true."),
        ("Khi nào scheduler đóng round?", "Khi `br.status = 'OPEN'`, `br.closesAt <= GETDATE()` và proposal còn `BOARD_REVIEW`."),
        ("Quorum là gì?", "Số phiếu hợp lệ tối thiểu để round được quyết định, đọc từ `proposal.minimumVoteQuorum`."),
        ("Thiếu quorum thì sao?", "System ghi history và mở vòng vote mới."),
        ("Hòa phiếu thì sao?", "System ghi note revote và mở vòng vote mới."),
        ("Approve thắng thì sao?", "Tạo `Series` nếu chưa có, proposal sang `APPROVED`, notify Mangaka và Tantou."),
        ("Revise thắng thì sao?", "Proposal sang `REVISION_REQUESTED`, notify Mangaka sửa."),
        ("Reject thắng thì sao?", "Proposal sang `REJECTED`, set rejected time."),
        ("Undo vote hoạt động thế nào?", "Board có thể xóa dòng vote trong `ProposalHistory` trong vòng 60 giây nếu proposal vẫn `BOARD_REVIEW`."),
        ("Class nào trả thông tin undo?", "`BoardVoteUndoInfo` chứa `historyId`, `decision`, `remainingSeconds`."),
        ("Vì sao cần `lockProposal()`?", "Để tránh race condition khi nhiều request cùng submit/review/vote/resolve một proposal."),
        ("`WITH (UPDLOCK, ROWLOCK)` có ý nghĩa gì?", "SQL Server khóa dòng proposal đang xử lý để transaction khác không sửa đồng thời."),
        ("`selectColumns()` làm gì?", "Ghép các cột base của Proposal với các cột tổng hợp board voting nếu schema sẵn sàng."),
        ("Vì sao có `isBoardVotingSchemaReady()`?", "Để code chạy được cả khi database cũ chưa có bảng board round hoặc cột `boardRoundId`."),
        ("`LEGACY_BOARD_SELECT_COLUMNS` dùng khi nào?", "Khi schema board voting chưa sẵn sàng, trả các giá trị vote mặc định 0/null."),
        ("Trang list hiển thị gì?", "`proposal/list.jsp` hiển thị proposal, status, progress board voting và link View."),
        ("Trang detail hiển thị gì?", "`proposal/detail.jsp` hiển thị thông tin proposal, voting summary, voter list, form review/vote và history."),
        ("Trang create gửi endpoint nào?", "`POST /main/proposals/create`."),
        ("Trang edit gửi endpoint nào?", "`POST /main/proposals/{id}/edit`."),
        ("Form Tantou review gửi endpoint nào?", "`POST /main/proposals/{id}/review`."),
        ("Form Board vote gửi endpoint nào?", "`POST /main/proposals/{id}/board-vote`."),
        ("Download file đi qua endpoint nào?", "`GET /main/proposals/{id}/file`."),
        ("Vì sao download file phải gọi `getDetail()` trước?", "Để kiểm quyền xem proposal trước khi stream file."),
        ("File upload được lưu ở đâu?", "Dưới `/uploads/proposals` trong webroot, path relative được lưu vào `sampleFilePath`."),
        ("Giới hạn upload hiện tại là bao nhiêu?", "20 MB, hằng `SAMPLE_FILE_MAX_SIZE_BYTES`."),
        ("Tên file được sanitize thế nào?", "Chỉ giữ ký tự `A-Z`, `a-z`, `0-9`, dấu chấm, gạch dưới và gạch ngang; thêm timestamp để tránh trùng."),
        ("`ProposalApiController` khác web controller ở điểm nào?", "API trả `ApiResponse`, còn web controller trả JSP/redirect; API create nhận path file thay vì multipart upload."),
        ("RBAC URL nằm ở đâu?", "`AuthInterceptor.isAllowed()`."),
        ("`/main/settings` ai được vào?", "Chỉ `ADMIN`."),
        ("`/main/proposals` ai được vào?", "`ADMIN`, `MANGAKA`, `TANTOU_EDITOR`, `EDITORIAL_BOARD`."),
        ("Vì sao service vẫn check quyền dù interceptor đã check?", "Interceptor chỉ check theo URL/role chung; service check theo dữ liệu cụ thể như owner, assigned editor, status."),
        ("Setting `proposal.maxSubmitAttempts` dùng ở đâu?", "`ProposalService.submitProposal()`, `submitForTantouReview()`, view detail/edit và `SystemSettingRepository.ensureSubmitAttemptConstraint()`."),
        ("Setting `proposal.minimumVoteQuorum` dùng ở đâu?", "`openBoardVotingRound()` ghi note và `resolveClosedBoardRound()` quyết định thiếu quorum."),
        ("Default max submit attempts là bao nhiêu?", "2, trong `ProposalSettingsService.DEFAULT_MAX_SUBMIT_ATTEMPTS`."),
        ("Default minimum quorum là bao nhiêu?", "3, trong `ProposalSettingsService.DEFAULT_MINIMUM_VOTE_QUORUM`."),
        ("Admin chỉnh setting qua JSP nào?", "`web/WEB-INF/jsp/settings/proposals.jsp`."),
        ("Repository nào lưu setting?", "`SystemSettingRepository`."),
        ("Vì sao update max attempts phải sửa constraint DB?", "Để database cũng bảo vệ `submitAttemptCount` không vượt giới hạn cấu hình."),
        ("Nếu max attempts mới nhỏ hơn dữ liệu hiện có thì sao?", "`ensureSubmitAttemptConstraint()` ném lỗi vì không thể tạo constraint hợp lệ."),
        ("Scheduler close expired chạy khi nào?", "Mỗi 15 phút theo cron `0 */15 * * * *`."),
        ("Scheduler remind vote chạy khi nào?", "Mỗi giờ theo cron `0 0 * * * *`."),
        ("Scheduler overdue Tantou chạy khi nào?", "Mỗi giờ phút 5 theo cron `0 5 * * * *`."),
        ("SLA Tantou review là bao lâu?", "48 giờ từ `submittedAt`, nếu quá thì set `tantouReviewOverdue = 1`."),
        ("Thông báo vote sắp đóng gửi cho ai?", "Các voter trong `ProposalBoardRoundVoter` chưa có vote history, khi round còn <=24 giờ."),
        ("Khi Board approve, notification gửi cho ai?", "Mangaka và Tantou assigned editor nếu có."),
        ("Vì sao dùng transaction trong repository?", "Vì một hành động thường gồm nhiều SQL: update proposal, insert history, insert notification, có thể tạo round/series."),
        ("Nếu transaction không rollback khi lỗi thì sao?", "Dữ liệu có thể lệch: status đổi nhưng không có history, hoặc vote được ghi nhưng round không đóng."),
        ("`ProposalHistory.actionType` có những giá trị nào?", "`CREATED`, `UPDATED`, `SUBMITTED`, `RESUBMITTED`, `ASSIGNED_EDITOR`, `APPROVED`, `REVISE_REQUESTED`, `REJECTED`."),
        ("`ProposalHistory.actorRole = SYSTEM` dùng khi nào?", "Khi hệ thống tự ghi sự kiện như assign editor, mở round, resolve vote."),
        ("`boardRoundId` trong history giúp gì?", "Biết vote thuộc vòng nào, không lẫn với vote của vòng trước."),
        ("Vì sao có `submitAttemptNumber`?", "Để phân biệt lịch sử của lần submit đầu và các lần resubmit."),
        ("Khi proposal revise rồi resubmit, attempt tăng không?", "Có, `submitForTantouReview()` tăng `submitAttemptCount` thêm 1."),
        ("Khi tạo draft, attempt là bao nhiêu?", "0, và history `CREATED` cũng ghi attempt 0."),
        ("`canEditDraft` trên controller dựa vào gì?", "User là `MANGAKA`, owner, status `DRAFT` hoặc `REVISION_REQUESTED`, và chưa đạt max attempts."),
        ("Vì sao Board/Admin không xem được draft?", "Draft là bản riêng của Mangaka, chưa nộp vào quy trình review."),
        ("Nếu Board vote sau khi round hết hạn thì sao?", "`voteByEditorialBoard()` đóng round và ném lỗi báo voting window expired, vòng mới đã mở nếu cần."),
        ("Nếu Board vote trùng thì sao?", "Service/repository ném lỗi `You have already voted on this proposal`."),
        ("Nếu voter không nằm trong round thì sao?", "Repository ném lỗi `You are not eligible to vote in this voting round`."),
        ("`listBoardRoundVoters()` trả những thông tin gì?", "User id, full name, đã vote chưa, decision, note, và có conflict không."),
        ("Vì sao voter conflict vẫn có thể hiển thị trong list?", "Để UI minh bạch rằng người đó là Board member nhưng bị excluded do conflict."),
        ("Điểm yếu lớn nhất của thiết kế vote hiện tại là gì?", "Vote nằm trong `ProposalHistory`, làm audit và dữ liệu quyết định bị trộn chung."),
        ("Refactor nên làm đầu tiên là gì?", "Tách vote sang bảng/service riêng hoặc tách notification/series creation ra khỏi `ProposalRepository`."),
        ("Vì sao dùng JDBC thay vì ORM?", "Project là Java web app truyền thống, repository hiện dùng plain JDBC; giữ pattern giúp đồng nhất codebase."),
        ("Vì sao không nên đưa Maven/Gradle vào?", "Project hiện là NetBeans/Ant; thêm build system mới có thể làm lệch cấu trúc và tăng rủi ro."),
        ("Nếu cần debug query proposal thì xem file nào?", "`ProposalRepository.java`."),
        ("Nếu cần sửa UI vote thì xem file nào?", "`proposal/detail.jsp`, `proposal/list.jsp`, `web/assets/js/proposal.js`, `web/assets/css/proposal.css`."),
        ("Nếu cần sửa quyền truy cập proposal thì xem đâu?", "`AuthInterceptor`, `ProposalService.getDetail()`, các method service liên quan role."),
        ("Nếu cần sửa số ngày vote thì xem đâu?", "`ProposalRepository.openBoardVotingRound()` ở câu SQL `DATEADD(DAY, 3, GETDATE())`."),
        ("Nếu cần sửa undo window thì xem đâu?", "`ProposalService.getBoardVoteUndoInfo()` và `ProposalRepository.undoBoardVote()` với mốc `60_000L`."),
        ("Nếu cần thêm decision mới cho Board thì phải sửa gì?", "Service validation, actionType mapping, DB check constraint `CK_PH_actionType`, query count vote, JSP radio options và resolve logic."),
        ("Nếu cần thêm status mới thì phải sửa gì?", "`ProposalStatus`, DB `CK_Proposal_status`, service/repository transition, JSP hiển thị status."),
        ("Nếu proposal approved mà đã có series thì sao?", "`ensureSeriesExistsOnApprove()` không tạo trùng, trả series id hiện có."),
        ("Tại sao `Series.proposalId` cần foreign key?", "Để đảm bảo series phải xuất phát từ proposal hợp lệ và có thể truy vết ngược."),
        ("Dashboard dùng proposal để làm gì?", "`DashboardController` đếm proposal, active proposal và hiển thị active proposal hiện tại."),
        ("Notification link proposal được resolve thế nào?", "`NotificationRepository` trả `/main/proposals/{id}` hoặc `/main/proposals/{id}/vote` tùy loại notification."),
        ("Nếu database cũ thiếu bảng board round thì chuyện gì xảy ra?", "`isBoardVotingSchemaReady()` trả false, select dùng legacy columns để tránh crash phần đọc." ),
        ("Điều cần nhớ nhất khi bảo vệ module là gì?", "Proposal là workflow kiểm duyệt nhiều bước: Mangaka tạo và submit, Tantou review, Board vote theo round/quorum/deadline, approve thì tạo Series, mọi bước đều ghi history và kiểm quyền."),
    ]
    for i, (q, a) in enumerate(questions, 1):
        lines.append(f"{i}. **{q}**")
        lines.append(f"   - Đáp án mẫu: {a}")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table):
    tbl_pr = table._tbl.tblPr
    margins = OxmlElement("w:tblCellMar")
    for m in ("top", "bottom", "start", "end"):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), "80" if m in ("top", "bottom") else "120")
        node.set(qn("w:type"), "dxa")
        margins.append(node)
    tbl_pr.append(margins)


def add_markdown_to_docx(md_text):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color in [("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if name == "Heading 2" else 18)
        style.paragraph_format.space_after = Pt(7 if name == "Heading 2" else 10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Tài Liệu Ôn Tập Module Proposal/Voting")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(12)

    in_code = False
    code_lines = []
    md_lines = md_text.splitlines()[1:]
    i = 0
    while i < len(md_lines):
        line = md_lines[i]
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                r = p.add_run("\n".join(code_lines))
                r.font.name = "Consolas"
                r.font.size = Pt(9)
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        if line.startswith("| "):
            table_lines = []
            while i < len(md_lines) and md_lines[i].startswith("| "):
                table_lines.append(md_lines[i])
                i += 1
            if len(table_lines) >= 2:
                rows = []
                for raw in table_lines:
                    cells = [c.strip().replace("<br>", "\n") for c in raw.strip().strip("|").split("|")]
                    if all(c == "---" for c in cells):
                        continue
                    rows.append(cells)
                if rows:
                    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    t.alignment = WD_TABLE_ALIGNMENT.CENTER
                    t.style = "Table Grid"
                    set_cell_margins(t)
                    for r_idx, row in enumerate(rows):
                        for c_idx, value in enumerate(row):
                            cell = t.cell(r_idx, c_idx)
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            cell.text = value.replace("`", "")
                            if r_idx == 0:
                                set_cell_shading(cell, "E8EEF5")
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    doc.add_paragraph()
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].replace("`", ""), style="List Bullet")
        elif line[:3].strip(".").isdigit() and ". **" in line:
            p = doc.add_paragraph(style="List Number")
            p.add_run(line.split("**", 2)[1]).bold = True
        elif line.startswith("   - Đáp án mẫu: "):
            doc.add_paragraph(line.replace("   - ", ""), style="List Bullet")
        else:
            doc.add_paragraph(line.replace("`", ""))
        i += 1
    return doc


def main():
    OUT_DIR.mkdir(exist_ok=True)
    md = build_markdown()
    MD_PATH.write_text(md, encoding="utf-8")
    doc = add_markdown_to_docx(md)
    doc.save(DOCX_PATH)
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
